"""Async file monitoring and processing system for Personal Knowledge Graph Server."""

import asyncio
import logging
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Set
from dataclasses import dataclass

# File processing imports
import PyPDF2
from docx import Document

# Watchdog imports
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler, FileCreatedEvent, FileModifiedEvent

# Local imports
from .config import get_config
from .cloud_nlp import process_file_content
from .knowledge_graph import store_extracted_data

logger = logging.getLogger(__name__)

@dataclass
class ProcessingResult:
    """Result of file processing operation."""
    file_path: str
    success: bool
    entities_count: int
    relationships_count: int
    processing_time: float
    error_message: Optional[str] = None

class FileEventHandler(FileSystemEventHandler):
    """Handle file system events for the monitored directory."""
    
    def __init__(self, file_monitor: 'FileMonitor'):
        super().__init__()
        self.file_monitor = file_monitor
        self.config = get_config()
        
        # Track recently created files to avoid duplicate processing
        self._recent_files: Set[str] = set()
    
    def on_created(self, event):
        """Handle file creation events."""
        if not event.is_directory:
            self._handle_file_event(event.src_path, "created")
    
    def on_modified(self, event):
        """Handle file modification events."""
        if not event.is_directory:
            self._handle_file_event(event.src_path, "modified")
    
    def _handle_file_event(self, file_path: str, event_type: str):
        """Handle file system events for supported file types."""
        path = Path(file_path)
        
        # Check if file extension is supported
        if path.suffix.lower() not in self.config.file_monitoring.supported_extensions:
            logger.debug(f"Ignoring unsupported file type: {path.suffix}")
            return
        
        # Check file size
        try:
            file_size = path.stat().st_size
            if file_size > self.config.file_monitoring.max_file_size:
                logger.warning(f"File too large ({file_size} bytes): {file_path}")
                return
            
            if file_size == 0:
                logger.debug(f"Empty file, skipping: {file_path}")
                return
        except OSError:
            logger.debug(f"Could not access file (may still be writing): {file_path}")
            return
        
        # Avoid duplicate processing
        file_key = f"{file_path}_{event_type}"
        if file_key in self._recent_files:
            logger.debug(f"Recently processed, skipping: {file_path}")
            return
        
        self._recent_files.add(file_key)
        
        logger.info(f"File {event_type}: {file_path}")
        
        # Schedule processing with delay to avoid partial file reads
        asyncio.create_task(self._delayed_process(file_path))
        
        # Clean up tracking set to prevent memory growth
        if len(self._recent_files) > 1000:
            self._recent_files.clear()
    
    async def _delayed_process(self, file_path: str):
        """Process file after a delay to ensure it's fully written."""
        try:
            # Wait for processing delay
            await asyncio.sleep(self.config.file_monitoring.processing_delay)
            
            # Check if file still exists and is stable
            path = Path(file_path)
            if not path.exists():
                logger.debug(f"File no longer exists: {file_path}")
                return
            
            # Check if file size is stable (hasn't changed in the last second)
            initial_size = path.stat().st_size
            await asyncio.sleep(1)
            
            if path.stat().st_size != initial_size:
                logger.debug(f"File still being written: {file_path}")
                # Retry once more
                await asyncio.sleep(2)
                if not path.exists():
                    return
            
            # Process the file
            await self.file_monitor.process_file(file_path)
            
        except Exception as e:
            logger.error(f"Error in delayed processing for {file_path}: {e}")

class FileMonitor:
    """Main file monitoring and processing class."""
    
    def __init__(self):
        """Initialize file monitor."""
        self.config = get_config()
        self.observer: Optional[Observer] = None
        self.event_handler: Optional[FileEventHandler] = None
        self._running = False
        self._processing_lock = asyncio.Lock()
        
        # Ensure directories exist
        self._ensure_directories()
        
        # Statistics
        self.stats = {
            "files_processed": 0,
            "files_failed": 0,
            "entities_extracted": 0,
            "relationships_extracted": 0,
            "total_processing_time": 0.0
        }
    
    def _ensure_directories(self):
        """Ensure required directories exist."""
        directories = [
            self.config.file_monitoring.watch_directory,
            self.config.file_monitoring.processed_directory,
            self.config.file_monitoring.inbox_directory
        ]
        
        for directory in directories:
            Path(directory).mkdir(parents=True, exist_ok=True)
            logger.debug(f"Ensured directory exists: {directory}")
    
    async def start_monitoring(self):
        """Start monitoring the watch directory for new files."""
        if self._running:
            logger.warning("File monitor is already running")
            return
        
        logger.info(f"Starting file monitor on: {self.config.file_monitoring.watch_directory}")
        
        # Set up watchdog observer
        self.event_handler = FileEventHandler(self)
        self.observer = Observer()
        self.observer.schedule(
            self.event_handler,
            self.config.file_monitoring.watch_directory,
            recursive=True
        )
        
        # Start observer
        self.observer.start()
        self._running = True
        
        logger.info("File monitor started successfully")
        
        try:
            # Process any existing files in inbox
            await self._process_existing_files()
            
            # Keep running until stopped
            while self._running:
                await asyncio.sleep(1)
                
        except Exception as e:
            logger.error(f"Error in file monitor: {e}")
        finally:
            await self.stop_monitoring()
    
    async def stop_monitoring(self):
        """Stop file monitoring."""
        if not self._running:
            return
        
        logger.info("Stopping file monitor")
        self._running = False
        
        if self.observer:
            self.observer.stop()
            self.observer.join()
        
        logger.info("File monitor stopped")
    
    async def _process_existing_files(self):
        """Process any files that already exist in the inbox directory."""
        inbox_path = Path(self.config.file_monitoring.inbox_directory)
        
        if not inbox_path.exists():
            return
        
        existing_files = []
        for extension in self.config.file_monitoring.supported_extensions:
            existing_files.extend(inbox_path.glob(f"*{extension}"))
        
        if existing_files:
            logger.info(f"Found {len(existing_files)} existing files to process")
            
            for file_path in existing_files:
                try:
                    await self.process_file(str(file_path))
                except Exception as e:
                    logger.error(f"Failed to process existing file {file_path}: {e}")
    
    async def process_file(self, file_path: str) -> ProcessingResult:
        """Process a single file through the AI pipeline.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            ProcessingResult with processing details
        """
        start_time = datetime.now()
        path = Path(file_path)
        
        # Use lock to prevent concurrent processing of the same file
        async with self._processing_lock:
            logger.info(f"Processing file: {file_path}")
            
            try:
                # Read file content
                content = self._read_file_content(file_path)
                
                if not content or not content.strip():
                    logger.warning(f"File is empty or unreadable: {file_path}")
                    return ProcessingResult(
                        file_path=file_path,
                        success=False,
                        entities_count=0,
                        relationships_count=0,
                        processing_time=0.0,
                        error_message="File is empty or unreadable"
                    )
                
                # Extract entities and relationships using AI
                entities, relationships = await process_file_content(
                    file_path=file_path,
                    content=content,
                    context=f"File: {path.name}"
                )
                
                # Store in knowledge graph
                entity_ids, relationship_ids = await store_extracted_data(
                    entities=entities,
                    relationships=relationships,
                    source_file=file_path
                )
                
                # Move file to processed directory
                await self._move_processed_file(file_path)
                
                # Calculate processing time
                processing_time = (datetime.now() - start_time).total_seconds()
                
                # Update statistics
                self.stats["files_processed"] += 1
                self.stats["entities_extracted"] += len(entities)
                self.stats["relationships_extracted"] += len(relationships)
                self.stats["total_processing_time"] += processing_time
                
                result = ProcessingResult(
                    file_path=file_path,
                    success=True,
                    entities_count=len(entities),
                    relationships_count=len(relationships),
                    processing_time=processing_time
                )
                
                logger.info(
                    f"Successfully processed {path.name}: "
                    f"{len(entities)} entities, {len(relationships)} relationships "
                    f"in {processing_time:.2f}s"
                )
                
                return result
                
            except Exception as e:
                processing_time = (datetime.now() - start_time).total_seconds()
                self.stats["files_failed"] += 1
                
                logger.error(f"Failed to process file {file_path}: {e}")
                
                return ProcessingResult(
                    file_path=file_path,
                    success=False,
                    entities_count=0,
                    relationships_count=0,
                    processing_time=processing_time,
                    error_message=str(e)
                )
    
    def _read_file_content(self, file_path: str) -> str:
        """Read content from various file types.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Text content of the file
            
        Raises:
            ValueError: If file type is not supported
            IOError: If file cannot be read
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        try:
            if extension in ['.md', '.txt']:
                return self._read_text_file(file_path)
            elif extension == '.pdf':
                return self._read_pdf_file(file_path)
            elif extension == '.docx':
                return self._read_docx_file(file_path)
            else:
                raise ValueError(f"Unsupported file type: {extension}")
                
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {e}")
            raise IOError(f"Could not read file {file_path}: {e}")
    
    def _read_text_file(self, file_path: str) -> str:
        """Read content from text/markdown files."""
        try:
            # Try UTF-8 first, fallback to other encodings
            encodings = ['utf-8', 'utf-8-sig', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise UnicodeDecodeError("Could not decode file with any supported encoding")
            
        except Exception as e:
            raise IOError(f"Error reading text file: {e}")
    
    def _read_pdf_file(self, file_path: str) -> str:
        """Read content from PDF files."""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(text)
                    except Exception as e:
                        logger.warning(f"Could not extract text from page {page_num} of {file_path}: {e}")
                        continue
            
            content = '\n\n'.join(text_content)
            
            if not content.strip():
                raise ValueError("No readable text found in PDF")
            
            return content
            
        except Exception as e:
            raise IOError(f"Error reading PDF file: {e}")
    
    def _read_docx_file(self, file_path: str) -> str:
        """Read content from Word documents."""
        try:
            doc = Document(file_path)
            
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            content = '\n\n'.join(text_content)
            
            if not content.strip():
                raise ValueError("No readable text found in DOCX file")
            
            return content
            
        except Exception as e:
            raise IOError(f"Error reading DOCX file: {e}")
    
    async def _move_processed_file(self, source_path: str):
        """Move processed file to the processed directory.
        
        Args:
            source_path: Path to the source file
        """
        try:
            source = Path(source_path)
            
            # Create processed directory structure
            processed_dir = Path(self.config.file_monitoring.processed_directory)
            
            # Add timestamp to filename to avoid conflicts
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            new_filename = f"{timestamp}_{source.name}"
            target_path = processed_dir / new_filename
            
            # Ensure target directory exists
            target_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Move the file
            shutil.move(str(source), str(target_path))
            
            logger.info(f"Moved processed file: {source.name} -> {target_path.name}")
            
        except Exception as e:
            logger.error(f"Failed to move processed file {source_path}: {e}")
            # Don't raise the error - processing was successful
    
    def get_statistics(self) -> Dict[str, any]:
        """Get processing statistics.
        
        Returns:
            Dictionary containing processing statistics
        """
        total_files = self.stats["files_processed"] + self.stats["files_failed"]
        avg_processing_time = (
            self.stats["total_processing_time"] / self.stats["files_processed"]
            if self.stats["files_processed"] > 0 else 0.0
        )
        
        return {
            "total_files": total_files,
            "files_processed": self.stats["files_processed"],
            "files_failed": self.stats["files_failed"],
            "success_rate": (
                self.stats["files_processed"] / total_files * 100
                if total_files > 0 else 0.0
            ),
            "entities_extracted": self.stats["entities_extracted"],
            "relationships_extracted": self.stats["relationships_extracted"],
            "total_processing_time": self.stats["total_processing_time"],
            "average_processing_time": avg_processing_time,
            "is_running": self._running
        }

# Convenience functions
async def start_file_monitoring():
    """Start file monitoring as a background task."""
    monitor = FileMonitor()
    await monitor.start_monitoring()

async def process_single_file(file_path: str) -> ProcessingResult:
    """Process a single file without starting the monitoring service.
    
    Args:
        file_path: Path to the file to process
        
    Returns:
        ProcessingResult with processing details
    """
    monitor = FileMonitor()
    return await monitor.process_file(file_path) 